"""
生成江汉区智能体大赛两份提交文档：
1. DemoGo_程序运行文档.docx
2. DemoGo_作品功能文档.docx
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = 'C:/Users/wei.gu/Documents/demogo/artifacts/'

# ─────────────────────────────────────────────────────────────
#  工具函数
# ─────────────────────────────────────────────────────────────

def add_heading(doc, text, level=1, color=None):
    style_name = f'Heading {level}'
    try:
        p = doc.add_heading(text, level=level)
    except Exception:
        p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_body(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    return p

def add_code(doc, text):
    """灰底等宽代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    # 设置灰色背景
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'F2F2F2')
    p._p.get_or_add_pPr().append(shading)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_num(doc, text, indent=False):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def set_page_margins(doc, top=2.5, bottom=2.5, left=3.0, right=3.0):
    from docx.shared import Cm
    section = doc.sections[0]
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)

# ─────────────────────────────────────────────────────────────
#  文档一：程序运行文档（部署说明）
# ─────────────────────────────────────────────────────────────

def build_deploy_doc():
    doc = Document()
    set_page_margins(doc)

    # 封面
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run('DemoGo 程序运行文档')
    tr.font.size = Pt(22)
    tr.bold = True
    tr.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run('试用链接生成与部署平台 · v0.9.35 · 2026年6月')
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()

    # 一、项目简介
    add_heading(doc, '一、项目简介', 1, color=(0x1A, 0x56, 0xDB))
    add_body(doc,
        'DemoGo（demogo.cn）是一个面向 AI 编程工具生态的"试用链接生成与部署平台"，'
        '核心功能为：接收用户上传的 AI 编程产物（静态网页、React/Vue/Vite 前端项目、'
        'Node.js 后端服务），自动完成项目识别、内容安全审查、构建/运行环境分配，'
        '30 秒内生成可对外分享的专属试用链接，解决国内 AI 产品"生成容易、体验难"的核心痛点。')

    add_body(doc, '作品已正式上线运行，访问地址：', bold=True)
    add_code(doc, '主站：https://demogo.cn')
    add_code(doc, 'API 健康检查：https://demogo.cn/api/health')
    add_code(doc, 'CLI 工具：npx @demogo-cn/cli@latest --version')

    add_body(doc, '作品提交形式：在线服务（SaaS）+ CLI 工具（npm 包）+ MCP 协议插件')

    # 二、环境依赖说明
    add_heading(doc, '二、环境依赖说明', 1, color=(0x1A, 0x56, 0xDB))

    add_heading(doc, '（一）服务器端运行环境', 2)
    add_bullet(doc, '操作系统：Linux（推荐 Ubuntu 22.04 LTS）')
    add_bullet(doc, 'Node.js：v20.x 及以上（LTS 版本）')
    add_bullet(doc, 'MySQL：8.0 及以上')
    add_bullet(doc, 'Nginx：1.24 及以上（用于反向代理和静态文件服务）')
    add_bullet(doc, 'systemd：用于进程守护管理')
    add_bullet(doc, '内存：2GB 及以上（推荐 4GB）')
    add_bullet(doc, '磁盘：系统盘 20GB + 数据盘（按需扩展）')
    add_bullet(doc, '端口：3001（后端服务）、80/443（Nginx）')

    add_heading(doc, '（二）客户端使用环境（终端用户）', 2)
    add_bullet(doc, '普通用户（Web 上传）：任意现代浏览器，无需安装任何软件')
    add_bullet(doc, 'CLI 用户：Node.js 16+ 环境，执行 npx @demogo-cn/cli@latest')
    add_bullet(doc, 'MCP 集成用户：支持 MCP 协议的 AI 编程工具（Cursor、Claude Code 等）')

    add_heading(doc, '（三）主要依赖库与版本', 2)
    add_bullet(doc, 'Express 4.x — HTTP 服务框架')
    add_bullet(doc, 'mysql2 3.x — MySQL 数据库驱动')
    add_bullet(doc, 'Bull 4.x / 自研队列 — 异步任务队列')
    add_bullet(doc, 'archiver / unzipper — 压缩包处理')
    add_bullet(doc, 'React 18 + Vite — 前端框架（管理端）')
    add_bullet(doc, 'Nodemailer — 邮件通知')
    add_bullet(doc, 'Umzug — 数据库迁移管理')

    # 三、访问与使用方式（在线服务）
    add_heading(doc, '三、访问与使用方式（在线服务）', 1, color=(0x1A, 0x56, 0xDB))

    add_heading(doc, '（一）Web 界面上传（推荐，无需安装）', 2)
    add_body(doc, '步骤如下：')
    add_num(doc, '打开浏览器，访问 https://demogo.cn')
    add_num(doc, '点击"注册"或"登录"，完成账号创建（邮箱 + 密码）')
    add_num(doc, '进入控制台，点击"新建 Demo"，将项目文件压缩为 .zip 格式后上传')
    add_num(doc, 'DemoGo 自动识别项目类型并开始部署，通常在 30 秒内完成')
    add_num(doc, '部署完成后，复制生成的专属链接，即可分享给任何用户试用')

    add_heading(doc, '（二）CLI 命令行上传', 2)
    add_body(doc, '适用于开发者在 AI 编程工具（Cursor/Claude Code 等）流程中直接集成：')
    add_code(doc, '# 安装 CLI（一次性）')
    add_code(doc, 'npm install -g @demogo-cn/cli')
    add_code(doc, '')
    add_code(doc, '# 在项目根目录执行上传')
    add_code(doc, 'demogo publish ./my-project.zip')
    add_code(doc, '')
    add_code(doc, '# 或使用 npx（无需安装）')
    add_code(doc, 'npx @demogo-cn/cli@latest publish ./my-project.zip')
    add_body(doc, '执行后，终端将输出专属试用链接，例如：')
    add_code(doc, '✓ 部署成功！试用链接：https://demo-xxxx.demogo.cn')

    add_heading(doc, '（三）MCP 协议集成（Cursor / Claude Code 等 AI 工具）', 2)
    add_body(doc, '在 AI 编程工具的 MCP 配置文件中添加 DemoGo MCP 服务器：')
    add_code(doc, '{')
    add_code(doc, '  "mcpServers": {')
    add_code(doc, '    "demogo": {')
    add_code(doc, '      "command": "npx",')
    add_code(doc, '      "args": ["demogo-mcp@latest"]')
    add_code(doc, '    }')
    add_code(doc, '  }')
    add_code(doc, '}')
    add_body(doc, '配置完成后，AI 工具可通过自然语言指令（"帮我发布这个项目"）自动调用 DemoGo 完成部署。')

    # 四、核心配置说明
    add_heading(doc, '四、核心配置说明（服务端）', 1, color=(0x1A, 0x56, 0xDB))
    add_body(doc, '服务端通过环境变量进行配置（.env 文件或 systemd Environment 指令）：')
    add_code(doc, '# 数据库配置')
    add_code(doc, 'DB_HOST=localhost')
    add_code(doc, 'DB_PORT=3306')
    add_code(doc, 'DB_NAME=demogo')
    add_code(doc, 'DB_USER=demogo_user')
    add_code(doc, 'DB_PASSWORD=<your_password>')
    add_code(doc, '')
    add_code(doc, '# 服务端口（默认 3001）')
    add_code(doc, 'PORT=3001')
    add_code(doc, '')
    add_code(doc, '# 公开访问基础 URL')
    add_code(doc, 'PUBLIC_BASE_URL=https://demogo.cn')
    add_code(doc, '')
    add_code(doc, '# 内容安全（可选，用于内容审查增强）')
    add_code(doc, 'CONTENT_SECURITY_KEY=<api_key_if_enabled>')
    add_code(doc, '')
    add_code(doc, '# 邮件通知（可选）')
    add_code(doc, 'SMTP_HOST=smtp.example.com')
    add_code(doc, 'SMTP_USER=notify@demogo.cn')
    add_code(doc, 'SMTP_PASS=<smtp_password>')

    # 五、服务启动与停止
    add_heading(doc, '五、服务启动与验证', 1, color=(0x1A, 0x56, 0xDB))

    add_heading(doc, '（一）检查服务状态', 2)
    add_code(doc, '# 查看服务运行状态')
    add_code(doc, 'systemctl status demogo-server')
    add_code(doc, '')
    add_code(doc, '# API 健康检查（返回 JSON 含版本号）')
    add_code(doc, 'curl https://demogo.cn/api/health')

    add_heading(doc, '（二）服务管理指令', 2)
    add_code(doc, '# 启动服务')
    add_code(doc, 'systemctl start demogo-server')
    add_code(doc, '')
    add_code(doc, '# 重启服务（部署新版本后执行）')
    add_code(doc, 'systemctl restart demogo-server')
    add_code(doc, '')
    add_code(doc, '# 停止服务')
    add_code(doc, 'systemctl stop demogo-server')
    add_code(doc, '')
    add_code(doc, '# 查看实时日志')
    add_code(doc, 'journalctl -u demogo-server -f')

    add_heading(doc, '（三）完整验证清单', 2)
    add_code(doc, '# 主站可访问')
    add_code(doc, 'curl -I https://demogo.cn              # 200 OK')
    add_code(doc, '')
    add_code(doc, '# API 健康检查')
    add_code(doc, 'curl https://demogo.cn/api/health      # {"version":"0.9.35",...}')
    add_code(doc, '')
    add_code(doc, '# 能力列表')
    add_code(doc, 'curl https://demogo.cn/api/hosting/capabilities')
    add_code(doc, '')
    add_code(doc, '# CLI 工具验证')
    add_code(doc, 'npx --yes @demogo-cn/cli@latest --version      # 0.9.35')
    add_code(doc, 'npx --yes @demogo-cn/cli@latest doctor --api https://demogo.cn')

    # 六、测试验证说明
    add_heading(doc, '六、测试验证说明', 1, color=(0x1A, 0x56, 0xDB))

    add_heading(doc, '（一）评审人员快速验证', 2)
    add_body(doc, '评审人员无需任何安装或部署，直接通过在线服务验证功能：')
    add_num(doc, '访问 https://demogo.cn，注册账号后登录')
    add_num(doc, '准备一个简单的测试文件：创建 index.html，写入 <h1>Hello DemoGo!</h1>，压缩为 test.zip')
    add_num(doc, '在控制台上传 test.zip，等待约 10-30 秒')
    add_num(doc, '点击生成的试用链接，浏览器打开后应显示"Hello DemoGo!"页面')
    add_num(doc, '验证完成：核心部署链路可用')

    add_heading(doc, '（二）进阶验证：Node.js 服务部署', 2)
    add_body(doc, '提供样例项目（联系参赛团队获取），包含 Express 服务器 + MySQL 数据库，上传后验证：')
    add_bullet(doc, '服务自动启动，API 接口可访问')
    add_bullet(doc, 'MySQL 数据库自动初始化（从 schema.sql）')
    add_bullet(doc, '数据库连接正常，增删改查可用')

    add_heading(doc, '（三）内置测试覆盖情况', 2)
    add_bullet(doc, '单元测试：136 项，全部通过')
    add_bullet(doc, '集成测试（E2E）：161 项，覆盖静态/构建/Node.js/MySQL 等 9 种场景')
    add_bullet(doc, '冒烟测试：关键链路验证，部署前强制执行')

    # 尾注
    doc.add_paragraph()
    note = doc.add_paragraph()
    note_run = note.add_run('如需进一步演示或协助，请联系参赛团队：古伟  |  demogo.cn')
    note_run.font.size = Pt(10)
    note_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER

    out_path = OUT_DIR + 'DemoGo_程序运行文档_v0.9.35.docx'
    doc.save(out_path)
    print(f'[OK] 程序运行文档已生成：{out_path}')
    return out_path


# ─────────────────────────────────────────────────────────────
#  文档二：作品功能文档
# ─────────────────────────────────────────────────────────────

def build_feature_doc():
    doc = Document()
    set_page_margins(doc)

    # 封面
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run('DemoGo')
    tr.font.size = Pt(28)
    tr.bold = True
    tr.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    sub1 = doc.add_paragraph()
    sub1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s1r = sub1.add_run('试用链接生成与部署平台')
    s1r.font.size = Pt(16)
    s1r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s2r = sub2.add_run('参赛赛道：AI + 垂直应用  |  版本：v0.9.35  |  2026年6月')
    s2r.font.size = Pt(11)
    s2r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()

    # 一、产品概述
    add_heading(doc, '一、产品概述', 1, color=(0x1A, 0x56, 0xDB))
    add_body(doc,
        'DemoGo 是面向 AI 编程工具生态的"试用链接生成与部署平台"，'
        '专为解决 AI 生成产品"最后一公里"问题而设计。')

    add_body(doc, '核心问题：', bold=True)
    add_body(doc,
        '2024年以来，Cursor、Claude Code、Windsurf 等 AI 编程工具爆炸式增长，每天有数百万个 AI '
        '生成的产品原型产生。但这些产品长期停留在开发者本地，原因是：'
        '普通用户不懂服务器部署；国内免费托管方案（Vercel、Netlify 等）访问不稳定；'
        '部署过程对非技术用户门槛极高。')

    add_body(doc, 'DemoGo 的解法：', bold=True)
    add_bullet(doc, '上传一个 .zip 压缩包，30 秒内自动完成部署')
    add_bullet(doc, '生成国内可稳定访问的专属试用链接')
    add_bullet(doc, '支持静态网页、前端构建项目、Node.js 后端服务、MySQL 数据库全链路')
    add_bullet(doc, '面向非技术用户设计，无需任何服务器知识')

    # 二、核心功能
    add_heading(doc, '二、核心功能详述', 1, color=(0x1A, 0x56, 0xDB))

    add_heading(doc, '（一）智能项目识别引擎', 2)
    add_body(doc,
        '用户上传压缩包后，DemoGo 自动分析项目结构，无需用户声明项目类型。'
        '支持识别的项目类型包括：')
    add_bullet(doc, '静态网站：根目录包含 index.html（直接 Nginx 静态服务）')
    add_bullet(doc, '前端构建产物：dist/build/out/public 目录（直接部署构建结果）')
    add_bullet(doc, 'React/Vue/Vite 源码项目：包含 package.json + src 目录（自动构建后部署）')
    add_bullet(doc, 'Node.js 后端服务：包含 server.js / app.js / index.js（容器化运行）')
    add_bullet(doc, 'Node.js + MySQL：包含 schema.sql（自动初始化数据库 + 注入连接配置）')
    add_bullet(doc, 'Supabase 外部后端：自动识别并注入 Supabase 连接参数')
    add_body(doc,
        '识别结果展示给用户确认，并给出"可以发布/无法发布/有风险"的明确结论，消除用户的不确定感。',
        indent=True)

    add_heading(doc, '（二）内容安全自动审查', 2)
    add_body(doc,
        '每次部署前自动扫描项目内容，检查：违禁关键词、恶意外链、敏感信息（API Key 泄露）、'
        '不合规内容。审查不通过则阻止发布并给出明确原因，保障平台内容安全合规。')

    add_heading(doc, '（三）全链路失败诊断', 2)
    add_body(doc,
        '构建或运行失败时，系统自动分析失败日志，给出中文诊断报告：'
        '"为什么失败、有什么影响、下一步怎么做"。'
        '覆盖 20+ 种常见失败场景（依赖缺失、端口冲突、内存不足、语法错误等），'
        '让非技术用户也能读懂错误原因。')

    add_heading(doc, '（四）持久化运行时管理', 2)
    add_body(doc,
        'Node.js 服务运行实例状态持久化存储，服务器重启后自动恢复运行中的 Demo 实例。'
        '内置运行实例数量限制（按套餐），防止资源滥用；运行超时自动回收，保障平台稳定性。')

    add_heading(doc, '（五）多渠道发布集成', 2)
    add_body(doc, 'DemoGo 提供 5 种发布方式，适配不同用户场景：')

    add_bullet(doc, 'Web 控制台（demogo.cn）：面向普通用户，图形界面上传和管理')
    add_bullet(doc, 'CLI 工具（@demogo-cn/cli）：面向开发者，命令行一键发布')
    add_bullet(doc, 'MCP 协议插件（demogo-mcp）：Cursor/Claude Code 等 AI 工具内直接调用')
    add_bullet(doc, 'Codex Plugin：OpenAI Codex 环境集成')
    add_bullet(doc, 'Agent API：RESTful API，支持任意 AI Agent 程序调用')

    add_heading(doc, '（六）自动表单数据收集', 2)
    add_body(doc,
        '检测项目中的 HTML 表单，自动接管表单提交，将用户填写的试用反馈数据汇总到 DemoGo 控制台，'
        '让产品发布者无需额外开发后端即可收集用户反馈数据。')

    # 三、技术架构
    add_heading(doc, '三、技术架构与自主研发说明', 1, color=(0x1A, 0x56, 0xDB))

    add_heading(doc, '（一）整体技术架构', 2)
    add_body(doc, '技术栈：')
    add_bullet(doc, '后端：Node.js + Express，服务入口 server.js，约 500 行精简版本')
    add_bullet(doc, '前端：React 18 + Vite + TypeScript，响应式管理控制台')
    add_bullet(doc, '数据库：MySQL 8.0 + Umzug 增量迁移管理')
    add_bullet(doc, '静态托管：Nginx 反向代理 + 二级域名路由')
    add_bullet(doc, '进程管理：systemd 守护 + 运行状态 JSON 持久化文件')
    add_bullet(doc, '异步任务：自研轻量级队列（Bull 可选），处理构建和部署任务')

    add_heading(doc, '（二）核心自主研发模块（23个服务模块）', 2)
    add_body(doc, '以下均为团队从零自研，不依赖第三方 AI 云服务实现核心功能：')
    add_bullet(doc, 'project-classifier-service — 项目类型识别引擎（基于文件树结构分析）')
    add_bullet(doc, 'build-service — 多类型项目构建编排（静态/npm build/Node.js）')
    add_bullet(doc, 'runtime-service — Node.js 运行环境管理（进程生命周期 + 持久化）')
    add_bullet(doc, 'deployment-executor — 统一部署执行器（同步/异步双路径共享代码）')
    add_bullet(doc, 'inspection-service — 项目检查引擎（内容安全 + 可部署性 + 诊断）')
    add_bullet(doc, 'demo-lifecycle-service — Demo 全生命周期管理')
    add_bullet(doc, 'failure-diagnosis-service — 失败诊断与中文报告生成')
    add_bullet(doc, 'archive-analyzer — 压缩包深度分析器（支持 zip/tar.gz）')
    add_bullet(doc, 'form-data-service — 表单数据自动收集与管理')
    add_bullet(doc, 'content-security-service — 内容安全审查（关键词 + 链接扫描）')
    add_bullet(doc, 'mysql-store — MySQL 数据库连接池 + 写锁并发控制')
    add_bullet(doc, 'notification-service — 邮件通知服务（部署完成/失败）')

    add_heading(doc, '（三）开源组件引用声明', 2)
    add_body(doc, '项目使用了以下主要开源组件（均为标准集成，核心业务逻辑完全自主研发）：')
    add_bullet(doc, 'Express 4.x — MIT License，用于 HTTP 路由框架')
    add_bullet(doc, 'React 18 — MIT License，用于前端 UI 框架')
    add_bullet(doc, 'Vite 5.x — MIT License，用于前端构建工具')
    add_bullet(doc, 'mysql2 — MIT License，用于 MySQL 数据库驱动')
    add_bullet(doc, 'Umzug — MIT License，用于数据库迁移管理')
    add_bullet(doc, 'archiver / unzipper — MIT License，用于压缩包处理')
    add_bullet(doc, 'nodemailer — MIT License，用于邮件发送')
    add_body(doc,
        '以上开源组件均为基础工具库，DemoGo 的核心价值（项目智能识别、安全审查、'
        '失败诊断、运行时管理、全渠道集成等）均为团队自主设计和实现。',
        indent=True)

    # 四、产品亮点与差异化
    add_heading(doc, '四、产品亮点与差异化优势', 1, color=(0x1A, 0x56, 0xDB))

    add_heading(doc, '亮点一：智能体全渠道接入', 2)
    add_body(doc,
        'DemoGo 是业内少数同时支持 MCP 协议（Model Context Protocol）和 Agent API 的部署平台。'
        'AI 编程工具（Cursor、Claude Code）的 AI 助手可以直接通过 MCP 调用 DemoGo，'
        '实现"AI 写代码 → AI 自动部署 → 生成试用链接"的全自动化流程，无需人工介入。')

    add_heading(doc, '亮点二：国内可稳定访问', 2)
    add_body(doc,
        '针对国内网络环境深度优化，自建域名 + Nginx 服务架构，'
        '对标 Vercel/Railway 的功能但专为国内用户场景设计，'
        '解决国际部署平台在中国大陆访问不稳定的核心痛点。')

    add_heading(doc, '亮点三：非技术用户友好', 2)
    add_body(doc,
        '产品原则是"把方便留给用户，把麻烦留给 DemoGo"。'
        '用户界面全程使用非技术语言，所有错误提示给出中文解释和操作建议，'
        '不需要用户理解服务器、Docker、端口等技术概念。')

    add_heading(doc, '亮点四：全链路自动化', 2)
    add_body(doc,
        '从压缩包上传到试用链接生成，全程零配置：'
        '项目类型自动识别、依赖自动安装、构建自动执行、环境变量自动注入、'
        '数据库自动初始化、子域名自动分配。')

    # 五、商业模式
    add_heading(doc, '五、商业化落地与应用场景', 1, color=(0x1A, 0x56, 0xDB))

    add_heading(doc, '（一）目标用户', 2)
    add_bullet(doc, 'AI 编程工具重度用户（产品经理、设计师、创业者）— 用 AI 工具做产品，需要快速分享验证')
    add_bullet(doc, '独立开发者 — 快速展示 Side Project 原型，收集用户反馈')
    add_bullet(doc, 'AI 培训机构 — 学员作业和项目演示的托管平台')
    add_bullet(doc, '企业内部工具团队 — 内部工具快速部署分发')

    add_heading(doc, '（二）商业模式', 2)
    add_bullet(doc, 'Freemium 免费套餐：基础功能免费，限制 Demo 数量和运行时长')
    add_bullet(doc, '个人专业版：99元/月，无限 Demo，Node.js 服务支持')
    add_bullet(doc, '团队版：499元/月，支持团队协作、自定义域名')
    add_bullet(doc, '按量付费：超出套餐部分按流量和运行时长计费')

    add_heading(doc, '（三）当前运营状态', 2)
    add_bullet(doc, '平台已上线运行，域名 demogo.cn 可正常访问')
    add_bullet(doc, 'CLI 工具已发布至 npm（@demogo-cn/cli），可通过 npx 直接使用')
    add_bullet(doc, 'MCP 插件已发布至 npm（demogo-mcp），可集成至 Cursor 等工具')
    add_bullet(doc, '正处于商业化验证阶段，持续迭代用户反馈')

    # 六、测试数据
    add_heading(doc, '六、技术指标与测试数据', 1, color=(0x1A, 0x56, 0xDB))

    add_body(doc, '代码规模（v0.9.35）：', bold=True)
    add_bullet(doc, '后端 JS 代码：约 8,500 行（23个服务模块 + 10个路由 + 4个中间件）')
    add_bullet(doc, '前端 TSX/TS 代码：约 6,200 行（React 管理控制台）')
    add_bullet(doc, '测试代码：约 1,800 行（136 单元测试 + 161 E2E 测试）')

    add_body(doc, '测试覆盖（v0.9.35 部署前门禁）：', bold=True)
    add_bullet(doc, '单元测试：136/136 通过（✓）')
    add_bullet(doc, 'E2E 集成测试：161/161 通过（✓）')
    add_bullet(doc, '冒烟测试：所有关键链路通过（✓）')
    add_bullet(doc, 'npm 安全审计：0 漏洞（✓）')

    add_body(doc, '性能指标（实测）：', bold=True)
    add_bullet(doc, '静态网站部署：平均 8-15 秒')
    add_bullet(doc, 'React/Vite 源码构建部署：平均 25-60 秒（含 npm install + build）')
    add_bullet(doc, 'Node.js 服务启动：平均 10-20 秒')
    add_bullet(doc, 'API 响应时间：P95 < 200ms')

    # 尾注
    doc.add_paragraph()
    doc.add_paragraph()
    end_p = doc.add_paragraph()
    end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_r = end_p.add_run('DemoGo — 让 AI 生成的产品立刻被用户用到')
    end_r.bold = True
    end_r.font.size = Pt(13)
    end_r.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = contact_p.add_run('参赛团队：古伟（DemoGo）  |  官网：https://demogo.cn  |  2026年6月')
    cr.font.size = Pt(10)
    cr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    out_path = OUT_DIR + 'DemoGo_作品功能文档_v0.9.35.docx'
    doc.save(out_path)
    print(f'[OK] 作品功能文档已生成：{out_path}')
    return out_path


if __name__ == '__main__':
    p1 = build_deploy_doc()
    p2 = build_feature_doc()
    print(f'\n两份文档均已生成：\n  1. {p1}\n  2. {p2}')
